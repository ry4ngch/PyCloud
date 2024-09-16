import tkinter as tk
from tkinter import filedialog, messagebox
import sympy as sp
from sympy import solve, symbols, init_printing
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import warnings
import math
import customtkinter as ctk
init_printing(use_unicode=True)

warnings.filterwarnings("ignore", category=UserWarning, module='tkinter')
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

class HydraulicCalculation:
    def __init__(self, diameter, length, flow_rate, roughness, viscosity, density=1000):
        self.diameter = diameter
        self.length = length
        self.flow_rate = flow_rate
        self.roughness = roughness
        self.viscosity = viscosity
        self.density = density
        self.gravity = 9.81

        # Define symbols
        self.D, self.L, self.Q, self.e, self.mu, self.f, self.h, self.g = symbols('D,L,Q,e,mu,f,h,g')

        # Define the Darcy-Weisbach formula (head loss)
        self.pressure_drop_formula = (self.f * self.L / self.D) * ((self.Q ** 2) / (sp.pi ** 2 * self.D ** 4 * self.g))
        #self.pressure_drop_formula = self.f * (self.L / self.D) * (self.ρ * self.v ** 2 / 2)
        self.eq = sp.Eq(self.h, self.pressure_drop_formula)

    def calculate_velocity(self):
        """Calculate flow velocity based on the flow rate and diameter."""
        return (4 * self.flow_rate) / (sp.pi * self.diameter ** 2)

    def calculate_reynolds_number(self):
        """Calculate the Reynolds number based on velocity, diameter, and viscosity."""
        velocity = self.calculate_velocity()
        reynolds_number = (self.density * velocity * self.diameter) / self.viscosity
        return reynolds_number

    def calculate_friction_factor(self):
        """Determine friction factor based on Reynolds number and pipe roughness."""
        reynolds_number = self.calculate_reynolds_number()

        if reynolds_number < 2300:
            # Laminar flow
            friction_factor = 64 / reynolds_number
        else:
            # Turbulent flow, use Colebrook-White equation
            def colebrook(f):
                return (-2 * math.log10((self.roughness / (3.7 * self.diameter)) +
                                        (2.51 / (reynolds_number * math.sqrt(f))))) - (1 / math.sqrt(f))

            # Use numerical method to solve for f (initial guess 0.02)
            friction_factor = self.newton_raphson(colebrook, 0.02)

        return friction_factor

    @staticmethod
    def newton_raphson(func, guess, tolerance=1e-6, max_iterations=100):
        """Numerically solve using Newton-Raphson method."""
        x = guess
        for _ in range(max_iterations):
            f_value = func(x)
            f_prime = (func(x + 1e-6) - f_value) / 1e-6  # Derivative using small increment
            if f_prime == 0:
                break  # Prevent division by zero
            new_x = x - f_value / f_prime
            if abs(new_x - x) < tolerance:
                return new_x
            x = new_x
        return x  # Return best guess if max_iterations is reached

    def calculate_pressure_drop(self):
        """Calculate the head loss using Darcy-Weisbach equation."""
        try:
            friction_factor = self.calculate_friction_factor()

            # Substitute user inputs and solve for head loss 'h'
            pressure_drop_equation = self.eq.subs({
                self.D: self.diameter,
                self.L: self.length,
                self.Q: self.flow_rate,
                self.f: friction_factor,
                self.e: self.roughness,
                self.g: self.gravity
            })
            h_value = solve(pressure_drop_equation, self.h)

            # Check if the solution exists and return it
            if h_value:
                return float(h_value[0]), friction_factor
            else:
                raise ValueError("Unable to solve for head loss.")
        except Exception as e:
            print(f"Error in calculating pressure drop: {e}")
            return None, None

    def get_report_content(self, view_eq_pretty: bool = False):
        pressure_drop, friction_factor = self.calculate_pressure_drop()
        if pressure_drop is None or friction_factor is None:
            return ["Error calculating results. Please check input values."]

        if view_eq_pretty:
            formula_str = self.get_equation()
        else:
            formula_str = str(self.pressure_drop_formula).replace('**', '^')

        # Convert SymPy expressions to numeric values
        reynolds_number = self.calculate_reynolds_number().evalf()
        velocity = self.calculate_velocity().evalf()

        # Convert friction_factor to float if it's a float already
        if isinstance(friction_factor, sp.Basic):
            friction_factor = friction_factor.evalf()
        friction_factor = float(friction_factor)  # Ensure it's a float

        content = [
            f"Diameter: {self.diameter} m",
            f"Length: {self.length} m",
            f"Flow Rate: {self.flow_rate} m^3/s",
            f"Roughness: {self.roughness} m",
            f"Viscosity: {self.viscosity} Pa.s",
            f"Density: {self.density} kg/m³",
            f"Velocity: {velocity:.2f} m/s",
            f"Reynolds Number: {reynolds_number:.2f}",
            f"Friction Factor: {friction_factor:.5f}",
            f"",
            f"---------------------------------",
            f"Pressure Drop (ΔP) Calculation:",
            f"ΔP = {formula_str}",
            f"ΔP = {pressure_drop:.2f} Pa"
        ]
        return content

    def get_equation(self):
        # Convert sympy equation to string
        return sp.pretty(self.pressure_drop_formula, use_unicode=True)


class FileHandler:
    @staticmethod
    def _generate_common_content(calculation, header=None):
        # If header is provided, add it to the content
        content = [header] if header else []
        content += calculation.get_report_content()
        return content

    @staticmethod
    def generate_report_txt(filename, calculation, header=None):
        try:
            content = FileHandler._generate_common_content(calculation, header)
            with open(filename, 'w') as file:
                for line in content:
                    file.write(line + '\n')
            print(f"TXT report generated: {filename}")
        except Exception as e:
            print(f"Error writing TXT file: {e}")

    @staticmethod
    def generate_report_docx(filename, calculation, header=None):
        try:
            doc = Document()
            if header:
                doc.add_heading(header, level=1)

            content = FileHandler._generate_common_content(calculation, header)
            for line in content:
                doc.add_paragraph(line)
            doc.save(filename)
            print(f"DOCX report generated: {filename}")
        except Exception as e:
            print(f"Error writing DOCX file: {e}")

    @staticmethod
    def generate_report_pdf(filename, calculation, header=None):
        try:
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            y = height - 100
            x = 100

            content = FileHandler._generate_common_content(calculation, header)

            c.setFont("Helvetica", 12)

            for line in content:
                if y < 100:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = height - 100

                c.drawString(x, y, line)
                y -= 20

            c.save()
            print(f"PDF report generated: {filename}")
        except Exception as e:
            print(f"Error writing PDF file: {e}")


class ReportGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title('Hydraulic Report Generator')
        self.report_header = 'Hydraulic Calculation Report'
        self.create_widgets()

    def create_widgets(self):
        ctk.CTkLabel(self.root, text="Diameter (m):").grid(row=0, column=0)
        self.diameter_entry = ctk.CTkEntry(self.root)
        self.diameter_entry.grid(row=0, column=1)

        ctk.CTkLabel(self.root, text="Length (m):").grid(row=1, column=0)
        self.length_entry = ctk.CTkEntry(self.root)
        self.length_entry.grid(row=1, column=1)

        ctk.CTkLabel(self.root, text="Flow Rate (m^3/s):").grid(row=2, column=0)
        self.flow_rate_entry = ctk.CTkEntry(self.root)
        self.flow_rate_entry.grid(row=2, column=1)

        ctk.CTkLabel(self.root, text="Roughness (m):").grid(row=3, column=0)
        self.roughness_entry = ctk.CTkEntry(self.root)
        self.roughness_entry.grid(row=3, column=1)

        ctk.CTkLabel(self.root, text="Dynamic Viscosity (Pa.s):").grid(row=4, column=0)
        self.viscosity_entry = ctk.CTkEntry(self.root)
        self.viscosity_entry.grid(row=4, column=1)

        ctk.CTkLabel(self.root, text="Density (kg/m^3):").grid(row=5, column=0)
        self.density_entry = ctk.CTkEntry(self.root)
        self.density_entry.grid(row=5, column=1)

        self.generate_button = ctk.CTkButton(self.root, text="Generate Report", command=self.generate_report)
        self.generate_button.grid(row=7, column=0, columnspan=1)

        self.view_button = ctk.CTkButton(self.root, text="View Results", command=self.view_report)
        self.view_button.grid(row=7, column=1, columnspan=1)

        ctk.CTkLabel(self.root, text="Select Format:").grid(row=8, column=0)
        self.format_var = tk.StringVar(self.root)
        self.format_var.set("TXT")
        self.format_menu = ctk.CTkOptionMenu(self.root, variable=self.format_var, values=["TXT", "DOCX", "PDF"])
        self.format_menu.grid(row=8, column=1)

        self.include_header_var = tk.IntVar()
        self.include_header_checkbox = ctk.CTkCheckBox(self.root, text="Include Header",
                                                      variable=self.include_header_var)
        self.include_header_checkbox.grid(row=9, column=0, columnspan=1)

        self.pretty_print_var = tk.IntVar()
        self.pretty_print_checkbox = ctk.CTkCheckBox(self.root, text="Pretty Print Equation",
                                                      variable=self.pretty_print_var)
        self.pretty_print_checkbox.grid(row=9, column=1, columnspan=1)

    def get_calc_result(self):
        try:
            diameter = float(self.diameter_entry.get())
            length = float(self.length_entry.get())
            flow_rate = float(self.flow_rate_entry.get())
            roughness = float(self.roughness_entry.get())
            viscosity = float(self.viscosity_entry.get())
            density = float(self.density_entry.get())
        except ValueError:
            messagebox.showerror("Input Error", "Please enter valid numeric values.")
            return None

        return HydraulicCalculation(diameter, length, flow_rate, roughness, viscosity, density)

    def generate_report(self):
        calc = self.get_calc_result()
        if calc is None:
            return

        format_type = self.format_var.get().lower()
        ext_map = {"txt": ".txt", "docx": ".docx", "pdf": ".pdf"}
        default_extension = ext_map.get(format_type)

        file_name = filedialog.asksaveasfilename(
            defaultextension=default_extension,
            title="Save Report As"
        )

        if not file_name:
            return

        header = self.report_header if self.include_header_var.get() else None

        if format_type == "txt":
            FileHandler.generate_report_txt(file_name, calc, header)
        elif format_type == "docx":
            FileHandler.generate_report_docx(file_name, calc, header)
        elif format_type == "pdf":
            FileHandler.generate_report_pdf(file_name, calc, header)
        else:
            messagebox.showerror("Error", "Unsupported file format selected")

    def view_report(self):
        calc = self.get_calc_result()
        if calc is None:
            return

        output_text = tk.Text(self.root, height=10, width=50)
        output_text.grid(row=10, column=0, columnspan=2)

        output_text.delete(1.0, tk.END)
        for line in calc.get_report_content(self.pretty_print_var.get()):
            output_text.insert(tk.END, line + '\n')


if __name__ == "__main__":
    root = ctk.CTk()
    app = ReportGeneratorApp(root)
    root.mainloop()